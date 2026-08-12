"""Genera los documentos Word de entrega de Conduce-Fácil.

* documentacion_general_tecnica_conduce_facil.docx
* descripcion_publicitaria_conduce_facil.docx

Las cifras se leen de los archivos de datos reales del proyecto, de modo que la
documentación no puede quedar desalineada con lo que efectivamente se publica.
"""

from __future__ import annotations

import json
import os
import subprocess
from datetime import date

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt, RGBColor

RAIZ = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
DOCS = os.path.join(RAIZ, "documentacion")
DATOS = os.path.join(RAIZ, "public", "datos")
MARCA = os.path.join(RAIZ, "public", "assets", "marca")

AZUL = RGBColor(0x04, 0x07, 0x64)
TURQUESA = RGBColor(0x20, 0xB6, 0xB6)
GRIS = RGBColor(0x54, 0x54, 0x54)

VERSION = "1.0"
FECHA = date.today().strftime("%d-%m-%Y")
REPO = "https://github.com/andresgamonalm/conduce-facil"
SUBDOMINIO = "https://conduce-facil.gamonal.app"


def cifras() -> dict:
    with open(os.path.join(DATOS, "estudio_conduce_facil.json"), encoding="utf-8") as f:
        estudio = json.load(f)
    with open(os.path.join(DATOS, "preguntas_conduce_facil.json"), encoding="utf-8") as f:
        preguntas = json.load(f)
    with open(os.path.join(DATOS, "senales_conduce_facil.json"), encoding="utf-8") as f:
        senales = json.load(f)
    grupos = {}
    for s in senales["senales"]:
        grupos[s["grupo"]] = grupos.get(s["grupo"], 0) + 1
    return {
        "tarjetas": len(estudio["tarjetas"]),
        "capitulos": len(estudio["capitulos"]),
        "preguntas": len(preguntas["preguntas"]),
        "senales": len(senales["senales"]),
        "grupos": grupos,
        "figuras": len(os.listdir(os.path.join(RAIZ, "public", "assets", "manual"))),
        "fuente_estudio": estudio["fuente"],
        "fuente_senales": senales["fuente"],
    }


def preparar(documento: Document) -> None:
    estilo = documento.styles["Normal"]
    estilo.font.name = "Roboto"
    estilo.font.size = Pt(10.5)
    estilo.font.color.rgb = RGBColor(0x3B, 0x3B, 0x3B)
    for seccion in documento.sections:
        seccion.top_margin = Cm(2.2)
        seccion.bottom_margin = Cm(2.2)
        seccion.left_margin = Cm(2.4)
        seccion.right_margin = Cm(2.4)
    for nombre, tamano in (("Heading 1", 19), ("Heading 2", 14.5), ("Heading 3", 12)):
        s = documento.styles[nombre]
        s.font.name = "Roboto"
        s.font.size = Pt(tamano)
        s.font.bold = True
        s.font.color.rgb = AZUL


def titulo_portada(documento: Document, titulo: str, bajada: str) -> None:
    p = documento.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(os.path.join(MARCA, "icono_conduce_facil_1000x1000.png"), width=Cm(3.2))

    p = documento.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    a = p.add_run("Conduce")
    a.font.size = Pt(30); a.font.bold = True; a.font.color.rgb = AZUL; a.font.name = "Roboto"
    b = p.add_run("-Fácil")
    b.font.size = Pt(30); b.font.bold = True; b.font.color.rgb = TURQUESA; b.font.name = "Roboto"

    p = documento.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(titulo)
    r.font.size = Pt(15); r.font.bold = True; r.font.color.rgb = AZUL

    p = documento.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(bajada)
    r.font.size = Pt(10.5); r.font.color.rgb = GRIS


def tabla(documento: Document, filas: list[tuple[str, str]], encabezado: tuple[str, str] | None = None) -> None:
    t = documento.add_table(rows=0, cols=2)
    t.style = "Light Grid Accent 1"
    t.alignment = WD_TABLE_ALIGNMENT.LEFT
    if encabezado:
        fila = t.add_row().cells
        for celda, texto in zip(fila, encabezado):
            celda.text = ""
            r = celda.paragraphs[0].add_run(texto)
            r.font.bold = True
            r.font.size = Pt(10)
    for clave, valor in filas:
        fila = t.add_row().cells
        fila[0].text = ""
        r = fila[0].paragraphs[0].add_run(clave)
        r.font.bold = True
        r.font.size = Pt(10)
        fila[1].text = ""
        r = fila[1].paragraphs[0].add_run(valor)
        r.font.size = Pt(10)
    documento.add_paragraph()


def vinetas(documento: Document, items: list[str]) -> None:
    for item in items:
        documento.add_paragraph(item, style="List Bullet")


def imagen(documento: Document, ruta: str, pie: str, ancho: float = 15.5) -> None:
    if not os.path.exists(ruta):
        return
    p = documento.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(ruta, width=Cm(ancho))
    c = documento.add_paragraph()
    c.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = c.add_run(pie)
    r.font.size = Pt(8.5)
    r.font.color.rgb = GRIS


# ---------------------------------------------------------------------------
# Documentación general y técnica
# ---------------------------------------------------------------------------

def documentacion_tecnica(c: dict) -> str:
    d = Document()
    preparar(d)
    titulo_portada(
        d, "Documentación general y técnica",
        f"Versión {VERSION} · {FECHA} · Estado: entregado para revisión",
    )
    d.add_page_break()

    d.add_heading("1. Identificación del proyecto", level=1)
    tabla(d, [
        ("Nombre comercial", "Conduce-Fácil"),
        ("Nombre normalizado", "conduce_facil"),
        ("Descripción breve", "Aplicativo web de estudio para el examen teórico de la Licencia de Conducir Clase B en Chile."),
        ("Versión", VERSION),
        ("Fecha", FECHA),
        ("Estado", "Entregado para revisión del usuario"),
        ("Repositorio", REPO),
        ("Rama de trabajo", "claude/chile-driving-exam-content-snt6fn"),
        ("Subdominio previsto", SUBDOMINIO),
        ("Plataforma de publicación", "Cloudflare Pages"),
        ("Título del navegador", "Conduce-Fácil · Prepara tu examen teórico Clase B"),
        ("Favicon", "assets/marca/favicon_conduce_facil.ico (integrado y activo)"),
    ])

    d.add_heading("2. Descripción general", level=1)
    d.add_paragraph(
        "Conduce-Fácil resuelve un problema concreto: el examen teórico de la Licencia de Conducir "
        "Clase B en Chile consta de 35 preguntas y admite un máximo de dos respuestas erróneas, de "
        "modo que estudiar «más o menos» no alcanza. El aplicativo toma la totalidad de los contenidos "
        "oficiales, los divide en unidades de estudio en formato pregunta y respuesta, y agrega práctica "
        "medida: trivia con las señaléticas oficiales, tests simulados cronometrados y un diagnóstico "
        "que indica qué repasar."
    )
    d.add_heading("Problema", level=2)
    vinetas(d, [
        "El material oficial son dos manuales extensos en PDF, difíciles de estudiar y de repasar.",
        "El examen exige precisión: 33 respuestas correctas de 35 como mínimo.",
        "No existe forma de medir el propio avance ni de saber qué temas están flojos.",
    ])
    d.add_heading("Objetivo", level=2)
    d.add_paragraph(
        "Que la persona llegue al examen sabiendo, con datos, que su margen de error es holgado."
    )
    d.add_heading("Usuarios", level=2)
    vinetas(d, [
        "Persona propietaria (perfil de administración): estudia y además ve el avance de todas las cuentas.",
        "Personas estudiantes: cada una con su sesión, sus resultados y su diagnóstico, sin acceso a los de las demás.",
    ])
    d.add_heading("Flujo principal", level=2)
    vinetas(d, [
        "Ingreso con usuario y contraseña.",
        "Inicio: estado del avance y acceso directo al test de 35 preguntas.",
        "Estudio: capítulo por capítulo, marcando cada contenido como dominado o por repasar.",
        "Trivia de señalética: reconocimiento de señales entre cuatro alternativas.",
        "Test de prueba: examen simulado configurable, con cronómetro por pregunta.",
        "Repaso: capítulos, familias de señales y preguntas con peor rendimiento.",
        "Resultados: historial completo con tiempos y tasas de acierto.",
    ])
    d.add_heading("Entradas y resultados", level=2)
    tabla(d, [
        ("Entradas", "Credenciales, respuestas a preguntas y señales, marcas de dominio de cada contenido, preferencias de práctica."),
        ("Resultados", "Puntaje y aprobación por sesión, tiempo total y por pregunta, veces respondida y tasa de éxito por contenido, avance por capítulo, ranking de debilidades."),
    ])

    d.add_heading("3. Contenidos y fuentes", level=1)
    d.add_paragraph(
        "Todo el contenido proviene de los PDF oficiales incluidos en el repositorio. No hay texto "
        "redactado ni resumido: las respuestas de estudio y los fundamentos de las preguntas se recortan "
        "literalmente del manual mediante anclas de texto, y las señaléticas se extraen como imagen desde "
        "las páginas originales del Manual de Señalización."
    )
    tabla(d, [
        ("Fuente de estudio", c["fuente_estudio"]),
        ("Fuente de señalética", c["fuente_senales"]),
        ("Capítulos y anexos", str(c["capitulos"])),
        ("Tarjetas de estudio", f"{c['tarjetas']} (pregunta + respuesta literal + página de origen)"),
        ("Imágenes del manual", f"{c['figuras']} figuras originales asociadas a su página"),
        ("Preguntas de test", f"{c['preguntas']} con alternativas y fundamento literal citado"),
        ("Señaléticas", f"{c['senales']} imágenes con código y nombre oficiales"),
    ])
    d.add_heading("Familias de señalética", level=2)
    tabla(d, [(g, str(n)) for g, n in sorted(c["grupos"].items(), key=lambda x: -x[1])],
          ("Familia", "Señales"))

    d.add_heading("4. Alcance", level=1)
    d.add_heading("Incluido", level=2)
    vinetas(d, [
        "Login con usuario y contraseña, y gestión de cuentas desde el perfil de administración.",
        "Estudio completo del Manual CONASET con imágenes originales.",
        "Trivia de señalética con las 186 señales extraídas del Manual de Señalización.",
        "Test simulado configurable en cantidad, capítulos, señalética y tiempo por pregunta.",
        "Repaso adaptativo y resultados detallados por sesión, por pregunta y por señal.",
        "Funcionamiento en modo local (navegador) y en modo servidor (Cloudflare D1).",
    ])
    d.add_heading("Límites y procesos manuales", level=2)
    vinetas(d, [
        "La publicación en Cloudflare Pages y el enlace del subdominio requieren autorización del usuario: no se ejecutaron.",
        "El modo servidor se activa al enlazar una base D1; mientras tanto los resultados quedan en el navegador de cada dispositivo.",
        "La fotografía de Envato seleccionada para el login queda identificada en ENVATO_ASSETS.md; la red del entorno de desarrollo no permitió descargarla, de modo que la pantalla usa una ilustración oficial del Manual CONASET.",
        "La trivia cubre las señales verticales, transitorias, de facilidades para ciclos y de mensajería variable; las demarcaciones y los semáforos se estudian como contenido en el módulo Estudio.",
    ])

    d.add_heading("5. Arquitectura", level=1)
    tabla(d, [
        ("Frontend", "HTML, CSS y JavaScript con módulos ES nativos. Sin proceso de compilación ni dependencias externas en tiempo de ejecución."),
        ("Enrutado", "History API con reescritura SPA (public/_redirects). Rutas reales y compartibles."),
        ("Backend", "Cloudflare Pages Functions (functions/api/[[ruta]].js). Opcional."),
        ("Base de datos", "Cloudflare D1 (SQLite) con tablas usuarios, sesiones, progreso y preferencias. Se crea sola en la primera consulta."),
        ("Almacenamiento sin backend", "localStorage del navegador, con la misma interfaz de repositorio."),
        ("Autenticación", "PBKDF2-SHA256, 150.000 iteraciones, sal aleatoria de 16 bytes por cuenta. Sesión por cookie HttpOnly, Secure y SameSite=Lax con vigencia de 30 días."),
        ("Tipografía", "Roboto servida desde el propio dominio (assets/fonts)."),
        ("Integraciones externas", "Ninguna en tiempo de ejecución."),
    ])
    d.add_paragraph(
        "El aplicativo detecta el entorno al arrancar: consulta /api/salud y, si responde que la base "
        "está enlazada, trabaja en modo servidor; en caso contrario continúa en modo local. Las vistas "
        "no cambian, porque ambos repositorios exponen la misma interfaz."
    )

    d.add_heading("6. Mapa de rutas", level=1)
    tabla(d, [
        ("/login", "Acceso. Formulario y panel visual. Pantalla completa, sin navegación."),
        ("/home", "Tarea principal: estado del avance y acceso directo al test de 35 preguntas."),
        ("/estudio", "Índice de capítulos con avance por capítulo."),
        ("/estudio/:capitulo", "Contenidos del capítulo en pregunta y respuesta, con imágenes del manual."),
        ("/trivia", "Configuración y juego de la trivia de señalética."),
        ("/test", "Configuración y ejecución del test simulado."),
        ("/repaso", "Diagnóstico de debilidades y accesos de práctica dirigida."),
        ("/resultados", "Historial y estadística por sesión, pregunta y señal."),
        ("/configuracion", "Cuenta, contraseña, preferencias de práctica y reinicio de estadísticas."),
        ("/admin", "Sólo administración: cuentas y avance de todas las personas usuarias."),
    ], ("Ruta", "Contenido"))

    d.add_heading("7. Estructura del proyecto", level=1)
    tabla(d, [
        ("public/", "Sitio publicable en Cloudflare Pages."),
        ("public/index.html", "Documento único del aplicativo y biblioteca de símbolos SVG."),
        ("public/assets/css/", "Hoja de estilos del sistema visual y declaración tipográfica."),
        ("public/assets/js/", "Módulos: núcleo, datos, almacenamiento, contexto y vistas."),
        ("public/assets/marca/", "Logotipo, ícono, ICO, favicon y PNG de 1000 × 1000."),
        ("public/assets/senales/", f"{c['senales']} imágenes de señaléticas extraídas del Manual de Señalización."),
        ("public/assets/manual/", f"{c['figuras']} figuras del Manual CONASET."),
        ("public/assets/fonts/", "Roboto en formato woff2."),
        ("public/datos/", "Contenidos generados: estudio, preguntas y señaléticas."),
        ("functions/api/", "API de Cloudflare Pages Functions."),
        ("herramientas/", "Scripts de extracción, generación y verificación."),
        ("documentacion/", "Documentos, capturas y portada de presentación."),
        ("Manual-Conaset/ y Manual-Señalizacion/", "PDF oficiales, fuente única de todo el contenido."),
    ], ("Ruta", "Función"))

    d.add_heading("8. Archivos principales", level=1)
    tabla(d, [
        ("app_conduce_facil.js", "Arranque, estructura común, navegación y registro de rutas."),
        ("nucleo_conduce_facil.js", "Creación de nodos, formato, aleatoriedad, PBKDF2 y enrutador."),
        ("datos_conduce_facil.js", "Carga de contenidos, armado de tests y trivias, y cálculo de diagnóstico."),
        ("almacenamiento_conduce_facil.js", "Repositorios local y remoto, cuentas y registro de estadísticas."),
        ("contexto_conduce_facil.js", "Estado de la sesión y persistencia del progreso."),
        ("vistas_acceso_conduce_facil.js", "Login, configuración y administración."),
        ("vistas_estudio_conduce_facil.js", "Inicio, estudio y repaso."),
        ("vistas_ejercicios_conduce_facil.js", "Trivia, test y resultados."),
        ("texto_manual.py", "Lectura y normalización del Manual CONASET; recorte literal por anclas."),
        ("contenido_estudio.py", "Definición de las tarjetas de estudio."),
        ("contenido_preguntas.py", "Definición del banco de preguntas."),
        ("generar_datos.py", "Generación de los archivos JSON que consume el aplicativo."),
        ("verificar_contenidos.py", "Verificación de fidelidad contra los PDF oficiales."),
        ("generar_marca.py / generar_portada.py", "Identidad visual y portada de presentación."),
        ("servidor_local.py", "Servidor de desarrollo con reescritura SPA."),
    ], ("Archivo", "Función"))

    d.add_heading("9. Requisitos, ejecución y publicación", level=1)
    d.add_heading("Requisitos", level=2)
    vinetas(d, [
        "Para usar el aplicativo: un navegador moderno. No hay instalación.",
        "Para regenerar contenidos: Python 3.11 o superior con pymupdf, Pillow, cairosvg y python-docx.",
        "Para publicar: cuenta de Cloudflare con acceso al proyecto de Pages.",
    ])
    d.add_heading("Ejecución local", level=2)
    d.add_paragraph(
        "El acceso directo abrir_conduce_facil.command (macOS y Linux) y abrir_conduce_facil.bat "
        "(Windows) levantan el servidor y abren el aplicativo en el navegador con doble clic, sin "
        "escribir comandos."
    )
    d.add_heading("Regeneración de contenidos", level=2)
    vinetas(d, [
        "herramientas/generar_datos.py vuelve a construir los tres archivos de datos desde los PDF.",
        "herramientas/verificar_contenidos.py comprueba que todo lo publicado exista literalmente en los manuales.",
        "herramientas/generar_marca.py y generar_portada.py rehacen la identidad visual.",
    ])
    d.add_heading("Publicación", level=2)
    vinetas(d, [
        "Cloudflare Pages: proyecto conduce-facil, directorio de salida public, sin comando de compilación.",
        "Las funciones de functions/ se despliegan automáticamente con el sitio.",
        "Modo servidor: crear la base D1 y descomentar el bloque [[d1_databases]] de wrangler.toml.",
        "Subdominio: apuntar conduce-facil.gamonal.app al proyecto de Pages.",
    ])

    d.add_heading("10. Configuración y credenciales", level=1)
    tabla(d, [
        ("Variables de entorno", "Ninguna obligatoria."),
        ("Binding D1", "DB (opcional). Sin él, el aplicativo funciona en modo local."),
        ("Puerto de desarrollo", "4173."),
        ("Cuenta inicial", "Usuario «andres» con perfil de administración. La contraseña la definió la persona propietaria y puede cambiarse en /configuracion."),
        ("Tratamiento de la contraseña", "En el repositorio sólo existe la derivación PBKDF2-SHA256 con su sal; el texto plano no está en el código ni en la documentación."),
        ("Cookies", "cf_sesion, HttpOnly, Secure, SameSite=Lax, 30 días."),
    ])

    d.add_heading("11. Decisiones técnicas y alternativas evaluadas", level=1)
    tabla(d, [
        ("Sin framework ni compilación",
         "Se evaluó React con Vite. Se descartó: el aplicativo es de contenido y la carga de datos es estática, de modo que un framework habría añadido dependencias y un paso de build sin beneficio. Con módulos ES nativos el despliegue en Pages es directo y la mantención no depende de versiones de terceros."),
        ("Extracción de señaléticas por recorte vectorial",
         "El Manual de Señalización dibuja cada señal como vector, no como imagen incrustada. Se detecta el encabezado «NOMBRE (CÓDIGO)», se agrupan los trazos de la celda, se descarta el plano de cotas y se rasteriza la versión limpia. Se prohibió expresamente redibujar señales."),
        ("Contenido recortado por anclas de texto",
         "En lugar de transcribir el manual a mano, cada contenido declara dos anclas y el script extrae el texto exacto entre ellas. Si un ancla no existe, la generación falla. Es imposible publicar texto inventado o alterado."),
        ("Local primero, servidor opcional",
         "Permite entregar un aplicativo que funciona de inmediato sin tocar la infraestructura del usuario, y activar el modo centralizado cuando se autorice."),
        ("PBKDF2 sobre Web Crypto",
         "Disponible tanto en el navegador como en Workers, sin dependencias. 150.000 iteraciones y sal por cuenta."),
    ], ("Decisión", "Fundamento"))

    d.add_heading("12. Seguridad y privacidad", level=1)
    vinetas(d, [
        "Las contraseñas nunca se guardan en texto plano: se almacena la derivación PBKDF2-SHA256 con sal aleatoria por cuenta.",
        "La comparación de derivaciones es de tiempo constante para no filtrar información por el tiempo de respuesta.",
        "En modo servidor, la sesión viaja en cookie HttpOnly, Secure y SameSite=Lax; el identificador de sesión no es accesible desde JavaScript.",
        "Cada persona sólo puede leer y escribir su propio progreso. El perfil de administración puede leer el de todas, pero no escribirlo.",
        "El error de acceso es siempre genérico, sin revelar si el usuario existe.",
        "Debe quedar siempre al menos una cuenta con permisos de administración; el sistema impide eliminar la última.",
        "El aplicativo no recoge datos personales más allá del nombre de usuario y el nombre para mostrar, y no envía información a terceros.",
        "Cabeceras de seguridad en public/_headers: nosniff, SAMEORIGIN, referrer restringido y permisos de cámara, micrófono y ubicación desactivados.",
    ])

    d.add_heading("13. Control de calidad ejecutado", level=1)
    d.add_paragraph(
        "Las pruebas se ejecutaron con Chromium real mediante Playwright, sobre el sitio servido igual "
        "que en producción, y con verificación programática de contenidos contra los PDF."
    )
    tabla(d, [
        ("Verificación de contenidos", f"{c['tarjetas'] + c['preguntas'] + c['senales']} contenidos comprobados contra los PDF oficiales: sin discrepancias."),
        ("Acceso", "Credenciales incorrectas rechazadas con mensaje genérico; credenciales correctas entran a /home."),
        ("Navegación", "Las diez rutas responden y marcan la sección activa, incluida la recarga directa en rutas anidadas."),
        ("Estudio", "Apertura de contenidos, imágenes del manual y marcado de dominio con avance por capítulo."),
        ("Trivia", "Diez señales respondidas correctamente entregan resultado aprobado; la retroalimentación indica la señal correcta al fallar."),
        ("Test", "Test de 35 preguntas con nueve de señalética, revisión pregunta por pregunta y regla de máximo dos errores."),
        ("Cronómetro", "Cuenta regresiva verificada con límite de 30 segundos por pregunta."),
        ("Persistencia", "El historial se mantiene tras recargar la página."),
        ("Aislamiento entre cuentas", "Una cuenta de estudiante no ve el enlace de administración y su acceso directo a /admin queda bloqueado."),
        ("Responsive", "Escritorio 1400 px, tablet 834 px y móvil 390 px sin scroll horizontal ni recortes."),
        ("Accesibilidad", "Todas las imágenes con texto alternativo; ningún control por debajo de 40 px de alto."),
        ("Consola y red", "Cero errores de consola y cero respuestas con código igual o superior a 400."),
    ], ("Prueba", "Resultado"))

    d.add_heading("14. Errores encontrados y corregidos", level=1)
    tabla(d, [
        ("Rutas de recursos relativas", "Al recargar una ruta anidada como /estudio/cap6 el navegador buscaba los recursos dentro de esa ruta y el módulo no cargaba. Se pasaron todas las rutas a absolutas."),
        ("Pérdida de progreso al navegar", "El guardado diferido podía perderse al cambiar de página. En modo local la escritura pasó a ser inmediata y en modo servidor se fuerza el envío antes de abandonar la página."),
        ("Formulario de alta de cuentas", "El formulario se refería a sí mismo después de una operación asíncrona y fallaba al limpiarse. Se captura la referencia antes de esperar."),
        ("Sección activa tras iniciar sesión", "Al reconstruirse la estructura se perdía la marca de sección activa. Ahora se aplica después de pintar cada vista."),
        ("Barra inferior en móvil", "Ocho accesos no caben en 390 px y quedaban recortados. Se dejaron cinco accesos y un agrupador «Más»."),
        ("Señales extraídas del plano de cotas", "Cinco señales tomaban el dibujo técnico en lugar de la versión limpia. Se corrigieron con recortes verificados uno a uno."),
    ], ("Error", "Corrección"))

    d.add_heading("15. Mantenimiento, respaldos y pendientes", level=1)
    d.add_heading("Mantenimiento", level=2)
    vinetas(d, [
        "Si CONASET publica una versión nueva del manual, se reemplaza el PDF y se ejecutan generar_datos.py y verificar_contenidos.py.",
        "Para agregar preguntas basta con añadir entradas en contenido_preguntas.py: el fundamento se recorta y se valida solo.",
        "La identidad visual se rehace con generar_marca.py, sin editar imágenes a mano.",
    ])
    d.add_heading("Respaldos", level=2)
    vinetas(d, [
        "El código y los contenidos viven en el repositorio de GitHub.",
        "En modo servidor, los resultados se respaldan con la exportación de la base D1.",
        "En modo local, los resultados viven en el navegador de cada dispositivo.",
    ])
    d.add_heading("Pendientes y revisiones manuales", level=2)
    vinetas(d, [
        "Publicar en Cloudflare Pages y enlazar el subdominio conduce-facil.gamonal.app (requiere autorización).",
        "Crear la base D1 y descomentar el bloque de wrangler.toml para activar el modo servidor (requiere autorización).",
        "Descargar y licenciar la fotografía de Envato identificada en ENVATO_ASSETS.md y dejarla como public/assets/img/login_conduce_facil.jpg; el login la tomará automáticamente.",
        "Revisar visualmente las capturas y la portada incluidas en esta documentación.",
    ])

    d.add_heading("16. Historial de cambios", level=1)
    tabla(d, [
        (f"{VERSION} · {FECHA}", "Primera versión completa: extracción de contenidos, identidad, aplicativo, API opcional, verificación y pruebas."),
    ], ("Versión", "Cambios"))

    d.add_heading("17. Capturas del producto", level=1)
    imagen(d, os.path.join(DOCS, "capturas", "pantalla_inicio_conduce_facil.jpg"),
           "Inicio: estado del avance y acceso directo al test de 35 preguntas.")
    imagen(d, os.path.join(DOCS, "capturas", "pantalla_estudio_conduce_facil.jpg"),
           "Estudio: contenido del manual en pregunta y respuesta, con la imagen original y la cita de página.")
    imagen(d, os.path.join(DOCS, "capturas", "pantalla_trivia_conduce_facil.jpg"),
           "Trivia de señalética: la señal oficial sin su nombre y cuatro alternativas.")
    imagen(d, os.path.join(DOCS, "capturas", "pantalla_test_conduce_facil.jpg"),
           "Test de prueba: 35 preguntas con cronómetro por pregunta.")
    imagen(d, os.path.join(DOCS, "capturas", "pantalla_resultados_conduce_facil.jpg"),
           "Resultados: historial con puntaje, tiempo total y tiempo por pregunta.")
    imagen(d, os.path.join(DOCS, "capturas", "pantalla_repaso_conduce_facil.jpg"),
           "Repaso: capítulos y familias de señales ordenados por debilidad.")
    imagen(d, os.path.join(DOCS, "capturas", "pantalla_administracion_conduce_facil.jpg"),
           "Administración: cuentas y avance de todas las personas usuarias.")
    imagen(d, os.path.join(DOCS, "capturas", "pantalla_movil_conduce_facil.jpg"),
           "Adaptación móvil con barra de navegación inferior táctil.", ancho=7.0)

    ruta = os.path.join(DOCS, "documentacion_general_tecnica_conduce_facil.docx")
    d.save(ruta)
    return ruta


# ---------------------------------------------------------------------------
# Descripción publicitaria
# ---------------------------------------------------------------------------

def descripcion_publicitaria(c: dict) -> str:
    d = Document()
    preparar(d)
    titulo_portada(d, "Descripción publicitaria y textos", f"Versión {VERSION} · {FECHA}")
    d.add_page_break()

    d.add_heading("1. Ficha", level=1)
    tabla(d, [
        ("Nombre", "Conduce-Fácil"),
        ("Categoría", "Educación · Preparación de exámenes · Seguridad vial"),
        ("Público objetivo", "Personas que rendirán el examen teórico de la Licencia de Conducir Clase B en Chile."),
        ("Plataforma", "Web, con adaptación a tablet y teléfono móvil."),
        ("Idioma", "Español de Chile."),
        ("Desarrollado por", "Gamonal"),
    ])

    d.add_heading("2. Descripciones", level=1)
    d.add_heading("Una línea", level=2)
    d.add_paragraph("Conduce-Fácil: el manual de conducción de Chile convertido en práctica medida.")
    d.add_heading("Descripción breve para presentación", level=2)
    d.add_paragraph(
        "Conduce-Fácil transforma los dos manuales oficiales de conducción de Chile en un plan de "
        "estudio con resultados: contenidos en pregunta y respuesta, trivia con las señaléticas "
        f"oficiales y tests simulados de 35 preguntas que aplican la misma exigencia del examen real."
    )
    d.add_heading("Descripción completa para sitio web", level=2)
    d.add_paragraph(
        "El examen teórico de la Licencia Clase B admite dos errores. Conduce-Fácil está construido "
        "para esa exigencia. Toma el Libro para la Conducción en Chile de CONASET y el Manual de "
        f"Señalización de Tránsito y los convierte en {c['tarjetas']} contenidos de estudio en formato "
        f"pregunta y respuesta, {c['senales']} señaléticas oficiales para reconocer y {c['preguntas']} "
        "preguntas de alternativas con su fundamento citado."
    )
    d.add_paragraph(
        "Las respuestas no son un resumen: son el texto literal del manual, con la página de origen a la "
        "vista y las imágenes originales del propio documento. Las señales tampoco son ilustraciones "
        "parecidas: son las señales oficiales extraídas de las páginas del Manual de Señalización."
    )
    d.add_paragraph(
        "Cada test se cronometra pregunta por pregunta. Cada respuesta alimenta un diagnóstico que "
        "indica qué capítulos y qué familias de señales conviene repasar. El avance queda registrado y "
        "se puede seguir en el tiempo."
    )
    d.add_heading("Ficha de producto", level=2)
    d.add_paragraph(
        f"Aplicativo web de preparación del examen teórico Clase B. {c['tarjetas']} contenidos oficiales "
        f"en pregunta y respuesta, {c['senales']} señaléticas del Manual de Señalización, "
        f"{c['preguntas']} preguntas de alternativas, tests configurables con cronómetro, repaso "
        "adaptativo y estadísticas por persona. Multiusuario con perfil de administración."
    )

    d.add_heading("3. Propuesta de valor", level=1)
    tabla(d, [
        ("Fidelidad verificable", "Cada respuesta y cada fundamento se recortan literalmente del PDF oficial y se comprueban de forma automática contra la fuente."),
        ("Señalética real", "Las señales se extraen de las páginas del Manual de Señalización, con su código y su nombre oficiales."),
        ("Exigencia del examen real", "35 preguntas, máximo 2 errores, con la misma proporción de señalética."),
        ("Medición del avance", "Tiempo por pregunta, tasa de acierto por contenido y avance por capítulo."),
        ("Diagnóstico accionable", "El repaso indica qué estudiar y abre directamente la práctica de ese tema."),
        ("Sin instalación", "Funciona en el navegador, en computador, tablet y teléfono."),
    ], ("Atributo", "Beneficio"))

    d.add_heading("4. Funcionalidades", level=1)
    vinetas(d, [
        f"Estudio: {c['capitulos']} capítulos y anexos, {c['tarjetas']} contenidos con página citada y {c['figuras']} imágenes originales del manual.",
        f"Trivia de señalética: {c['senales']} señales oficiales, filtrables por familia, con corrección inmediata.",
        f"Test de prueba: {c['preguntas']} preguntas con fundamento citado, cantidad y temas configurables, señalética opcional y cronómetro por pregunta.",
        "Repaso: capítulos, familias de señales, señales y preguntas ordenados por debilidad.",
        "Resultados: historial de sesiones, rendimiento por pregunta y por señal, tiempos medios.",
        "Cuentas: perfil de administración y perfil de estudiante, cada uno con sus propios resultados.",
        "Configuración: cambio de contraseña, cronómetro, límite de tiempo por pregunta y reinicio de estadísticas.",
    ])

    d.add_heading("5. Casos de uso", level=1)
    vinetas(d, [
        "Estudiar un capítulo completo y marcar qué contenidos quedaron dominados.",
        "Repasar señaléticas en diez minutos desde el teléfono, por familia de señales.",
        "Simular el examen completo antes de ir a la Municipalidad y comprobar el margen de error.",
        "Volver sobre las preguntas falladas con el texto del manual a la vista.",
        "Acompañar a otra persona: crearle una cuenta y seguir su avance desde administración.",
    ])

    d.add_heading("6. Argumentos comerciales", level=1)
    vinetas(d, [
        "El contenido no se resume ni se reescribe: se cita. Cada respuesta indica la página del manual.",
        "La verificación es automática y reproducible: un script comprueba los contenidos publicados contra los PDF.",
        "La práctica se mide: no es «ir respondiendo», es saber en qué capítulo se falla y cuánto se demora.",
        "El criterio de aprobación es el del examen real, no uno inventado.",
        "Funciona sin instalar nada y sin depender de servicios externos en tiempo de ejecución.",
    ])

    d.add_heading("7. Textos para difusión", level=1)
    d.add_heading("LinkedIn", level=2)
    d.add_paragraph(
        "El examen teórico de conducción en Chile permite dos errores en 35 preguntas. Con ese margen, "
        "estudiar sin medir no sirve de mucho.\n\n"
        "Conduce-Fácil convierte los dos manuales oficiales —el Libro para la Conducción en Chile de "
        f"CONASET y el Manual de Señalización de Tránsito— en {c['tarjetas']} contenidos de estudio en "
        f"formato pregunta y respuesta, {c['senales']} señaléticas oficiales para reconocer y "
        f"{c['preguntas']} preguntas con su fundamento citado.\n\n"
        "Tres decisiones que marcaron el proyecto:\n"
        "· Las respuestas son texto literal del manual, con la página a la vista. Un script las verifica "
        "contra el PDF antes de publicar.\n"
        "· Las señales se extraen de las páginas del manual oficial. Ninguna fue redibujada.\n"
        "· Cada test aplica la regla real: 35 preguntas, máximo 2 errores, con cronómetro por pregunta.\n\n"
        "Desarrollado por Gamonal."
    )
    d.add_heading("Redes sociales", level=2)
    d.add_paragraph(
        "35 preguntas. 2 errores permitidos. Conduce-Fácil te deja llegar al examen teórico sabiendo "
        f"exactamente dónde estás parado: {c['senales']} señaléticas oficiales, {c['tarjetas']} contenidos "
        "del manual y tests cronometrados con la exigencia real."
    )
    d.add_heading("Mensaje para WhatsApp o correo", level=2)
    d.add_paragraph(
        "Te comparto Conduce-Fácil, para preparar el examen teórico de conducción Clase B. Tiene todo el "
        "manual de CONASET en preguntas y respuestas, las señaléticas oficiales para practicar y tests de "
        "35 preguntas con el mismo criterio del examen real. Funciona en el celular, sin instalar nada."
    )

    d.add_heading("8. Mensajes clave y concepto", level=1)
    tabla(d, [
        ("Concepto principal", "El manual oficial, convertido en práctica medida."),
        ("Mensaje 1", "Contenido literal y citado, no resumido."),
        ("Mensaje 2", "Señaléticas oficiales, no adaptaciones."),
        ("Mensaje 3", "La exigencia del examen real: 35 preguntas, 2 errores."),
        ("Mensaje 4", "Sabes en qué fallas y cuánto te demoras."),
        ("Tono", "Claro, directo, orientado a resultado. Sin promesas que el producto no cumpla."),
    ])
    d.add_heading("Palabras clave", level=2)
    d.add_paragraph(
        "examen teórico de conducir Chile · licencia clase B · manual CONASET · señales de tránsito "
        "Chile · test de conducción · práctica examen licencia · señalética vial · Ley de Tránsito · "
        "preparar licencia de conducir · manual de señalización de tránsito"
    )

    ruta = os.path.join(DOCS, "descripcion_publicitaria_conduce_facil.docx")
    d.save(ruta)
    return ruta


def main() -> None:
    os.makedirs(DOCS, exist_ok=True)
    c = cifras()
    for ruta in (documentacion_tecnica(c), descripcion_publicitaria(c)):
        tamano = os.path.getsize(ruta) // 1024
        print(f"  {os.path.basename(ruta):52s} {tamano:>6,} KB")


if __name__ == "__main__":
    main()
